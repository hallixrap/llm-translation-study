"""
Generate the revised manuscript as a .docx file using python-docx.
Output: /Users/chuka/Library/CloudStorage/OneDrive-Stanford/AI translation/Evaluating_LLM_Translation_Fidelity_REVISED.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = (
    "/Users/chukanya/Library/CloudStorage/OneDrive-Stanford/AI translation/"
    "Evaluating_LLM_Translation_Fidelity_REVISED.docx"
)


def add_paragraph(doc, text, style="Normal", bold=False, italic=False, alignment=None):
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bold_label_paragraph(doc, label, text, style="Normal"):
    """Add a paragraph where the label is bold and the rest is normal."""
    p = doc.add_paragraph(style=style)
    run_label = p.add_run(label)
    run_label.bold = True
    run_text = p.add_run(text)
    run_text.bold = False
    return p


def add_table(doc, headers, rows, footer_note=None):
    """Add a formatted table with headers, rows, and optional footer note."""
    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 for header row
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_val in enumerate(row_data):
            row_cells[col_idx].text = str(cell_val)

    if footer_note:
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(footer_note)
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table
    return table


def build_manuscript():
    doc = Document()

    # -------------------------------------------------------------------------
    # TITLE PAGE
    # -------------------------------------------------------------------------
    add_heading(doc, "Evaluating Large Language Model Translation Fidelity for "
                "Medical Documents Across High- and Low-Resource Languages", level=1)

    add_paragraph(
        doc,
        "Chukwuebuka Anyaegbuna, MD; Eduardo Juan Perez Guerrero; Jerry Liu, MD; "
        "Timothy Keyes, PhD; April Liang, MD; Natasha Steele, MD, MPH; "
        "Stephen Ma, MD; Jonathan Chen, MD, PhD; Kevin Schulman, MD, MBA",
        bold=True,
    )

    add_paragraph(doc, "")

    add_bold_label_paragraph(
        doc,
        "Corresponding Author: ",
        "Chukwuebuka Anyaegbuna, MD, 3180 Porter Drive, Palo Alto, CA 94304; "
        "gozirim@gmail.com; +1 628 309 3402",
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # ABSTRACT
    # -------------------------------------------------------------------------
    add_heading(doc, "Abstract", level=1)

    abstract_sections = [
        (
            "Importance: ",
            "For the approximately 27.3 million U.S. residents who use languages other than "
            "English at home, language barriers create substantial obstacles to accessing medical "
            "information and achieving optimal health outcomes. Professional medical translation "
            "remains costly and often unavailable, particularly for low-resource languages with "
            "limited digital representation.",
        ),
        (
            "Objective: ",
            "To evaluate whether frontier large language models (LLMs) maintain translation "
            "fidelity for medical documents across high-resource and low-resource languages, and "
            "to assess LLM translation quality relative to professional human translation as a "
            "co-primary outcome.",
        ),
        (
            "Design, Setting, and Participants: ",
            "Cross-sectional evaluation of four frontier LLMs (GPT-5.1, Claude Opus 4.5, "
            "Gemini 3 Pro, Kimi K2) translating 22 professionally-translated medical documents "
            "from CDC Vaccine Information Statements and American Cancer Society patient education "
            "materials into 8 languages classified by CommonCrawl representation: high-resource "
            "(Spanish, Chinese, Russian, Vietnamese), medium-resource (Korean, Arabic), and "
            "low-resource (Tagalog, Haitian Creole).",
        ),
        (
            "Main Outcomes and Measures: ",
            "Two co-primary outcomes: (1) back-translation fidelity assessed using LaBSE "
            "(semantic similarity, 0-1 scale) and BLEU (lexical overlap, 0-100 scale); and "
            "(2) comparison to professional translations using COMET and BERTScore. A cross-model "
            "back-translation sensitivity analysis was conducted to address potential same-model "
            "circularity: each model's forward translations were back-translated by all other "
            "models (1,576 additional back-translations across 9 model pairings).",
        ),
        (
            "Results: ",
            "Across 702 of 704 translation pairs (99.7%), all models achieved high semantic "
            "preservation (LaBSE greater than 0.92). Low-resource languages achieved semantic "
            "similarity that did not differ significantly from high-resource languages "
            "(Tagalog: 0.950; Haitian Creole: 0.955 vs Spanish: 0.954; p = 0.066), though this "
            "represents absence of detected difference rather than demonstrated equivalence. LLM "
            "translations approached professional quality on the co-primary comparison "
            "(COMET: 0.87-0.88 across models), with no significant differences between models on "
            "this measure (p = 0.56). In the cross-model sensitivity analysis, LaBSE scores were "
            "stable across all nine model pairings (overall same-model mean: 0.9497; cross-model "
            "mean: 0.9488; mean delta: -0.0009), with no model showing the substantial inflation "
            "expected if same-model circularity were driving results.",
        ),
        (
            "Conclusions and Relevance: ",
            "Frontier LLMs preserve medical meaning across diverse languages in standardized "
            "patient education materials, with performance approaching professional translation "
            "on automated metrics. The cross-model sensitivity analysis substantially mitigates "
            "concerns about same-model circularity. These findings warrant further evaluation "
            "with human comprehension outcomes and broader document types before informing "
            "clinical deployment decisions.",
        ),
    ]

    for label, text in abstract_sections:
        add_bold_label_paragraph(doc, label, text)

    doc.add_paragraph()
    add_paragraph(
        doc,
        "Keywords: machine translation, large language models, medical translation, health "
        "equity, language access, low-resource languages",
        italic=True,
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # KEY POINTS
    # -------------------------------------------------------------------------
    add_heading(doc, "Key Points", level=1)

    key_points = [
        (
            "Question: ",
            "Do large language models reliably translate medical documents for patients who use "
            "languages other than English, including speakers of low-resource languages with "
            "limited digital representation?",
        ),
        (
            "Findings: ",
            "In this cross-sectional evaluation of 702 translations across 8 languages, frontier "
            "LLMs preserved medical meaning with high fidelity (LaBSE greater than 0.92 for all "
            "models) and approached professional translation quality (COMET: 0.87-0.88). A "
            "cross-model sensitivity analysis (1,576 back-translations across 9 model pairings) "
            "found a mean difference of -0.0009 compared to same-model back-translation, "
            "substantially mitigating circularity concerns. Low-resource languages did not perform "
            "significantly worse than high-resource languages (p = 0.066), though this represents "
            "absence of detected difference rather than demonstrated equivalence.",
        ),
        (
            "Meaning: ",
            "Frontier LLMs show promise for extending medical translation access to historically "
            "underserved language communities, though validation with patient comprehension "
            "outcomes and broader document types is needed before informing clinical deployment "
            "or regulatory decisions.",
        ),
    ]

    for label, text in key_points:
        add_bold_label_paragraph(doc, label, text)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # INTRODUCTION
    # -------------------------------------------------------------------------
    add_heading(doc, "Introduction", level=1)

    intro_paras = [
        "Access to medical information is fundamental to patient engagement and health outcomes. "
        "Studies encompassing over 74,000 patients have demonstrated that patient education "
        "significantly improves physiological, physical, and psychological outcomes while reducing "
        "medication use and healthcare utilization.(1) Yet for the approximately 27.3 million "
        "U.S. residents who use languages other than English at home, language barriers create "
        "substantial obstacles to accessing this information.(2-4)",

        "The consequences are well-documented. Adults with non-English language preference report "
        "poorer health outcomes, including lower rates of glycemic control, higher rates of "
        "uncontrolled asthma, and increased odds of poorly controlled hypertension.(3) They are "
        "significantly less likely to utilize outpatient visits and are three times as likely as "
        "English-proficient patients to be uninsured.(3) Medical errors experienced by patients "
        "with non-English language preference are more likely to cause physical harm compared to "
        "those experienced by English-proficient patients.(4) As the Joint Commission notes in "
        "their Roadmap for Hospitals, \"effective communication is now accepted as an essential "
        "component of quality care and patient safety.\"(5)",

        "Federal law recognizes these disparities. Section 1557 of the Affordable Care Act and "
        "the National Standards for Culturally and Linguistically Appropriate Services (CLAS) "
        "mandate that healthcare organizations that receive federal funding provide language "
        "access services.(6,7) However, manual professional medical translation is costly and "
        "resource-intensive and is not always available in a timely or workflow-compatible manner "
        "during routine clinical care.(8,9) As a result, professional translation is often "
        "applied selectively rather than comprehensively, with documented gaps for speakers of "
        "less common languages and for written materials generated during rapid clinical processes "
        "such as hospital discharge.(10)",

        "Large language models have demonstrated substantial capabilities in translation "
        "tasks.(11) Given their ready accessibility, these technologies could help address gaps "
        "in medical information availability. However, medical translation presents unique "
        "challenges: terminology must be precise, instructions unambiguous, and errors can have "
        "serious consequences.(12) Early evaluations suggest LLM translation quality varies by "
        "language. A recent study found ChatGPT and Google Translate performed comparably to "
        "professional translation for Spanish and Portuguese discharge instructions but had "
        "significantly more clinically significant errors for Haitian Creole.(13) The reliability "
        "of LLMs across languages, particularly low-resource languages with limited training "
        "data, requires systematic evaluation.",

        "Evaluating this technology is an ongoing need given its rapid evolution. Recent survey "
        "data show that 57% of U.S. physicians are already using or planning to adopt AI "
        "translation services within the next year, a faster adoption rate than any other AI use "
        "case surveyed.(14) This creates urgency: clinicians are deploying tools with variable "
        "performance across languages, yet lack systematic data to guide their use.",

        "Traditional translation evaluation requires a bilingual human expert for each language "
        "pair, limiting scalability. Recent advances in multilingual natural language processing "
        "(NLP) have produced automated metrics (including LaBSE, COMET, and BERTScore) that "
        "capture semantic similarity across languages.(15-17) These metrics enable evaluation "
        "without requiring human evaluators for each language pair, making them particularly "
        "valuable for assessing low-resource languages where expert evaluators are scarce. We "
        "note that automated metrics are not equivalent to human evaluation, particularly for "
        "assessing clinical accuracy or patient comprehension, and we interpret them accordingly "
        "as a necessary complement to, rather than replacement for, human judgment.",

        "This study reports a systematic assessment of LLM translation performance using "
        "automated multilingual metrics across eight languages, with two co-primary outcomes: "
        "back-translation fidelity and comparison to professional translation. We focus on "
        "English-to-target translation, reflecting clinical scenarios where English-language "
        "materials must be made accessible to patients with non-English language preference. We "
        "employed back-translation validation (translating text to the target language and back "
        "to confirm meaning preservation) as a consistency check, grounded in cross-cultural "
        "adaptation frameworks.(18,19) We acknowledge that back-translation is not a standard "
        "methodology in linguistic scholarship as a standalone validation tool, and we address "
        "its limitations directly, including a cross-model sensitivity analysis that replaces "
        "the back-translating model to test for same-model circularity.",
    ]

    for para in intro_paras:
        add_paragraph(doc, para)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # METHODS
    # -------------------------------------------------------------------------
    add_heading(doc, "Methods", level=1)

    add_heading(doc, "Study Objectives", level=2)
    add_paragraph(
        doc,
        "This study addressed three research questions: (1) Do LLM translations preserve medical "
        "meaning across high-resource and low-resource languages? (2) Does translation fidelity "
        "differ by language resource level? (3) How do LLM translations compare to professional "
        "translations? The third question was designated as a co-primary outcome rather than a "
        "secondary validation.",
    )
    add_paragraph(
        doc,
        "We hypothesized that current-generation LLMs maintain translation fidelity across both "
        "high- and low-resource languages, achieving semantic preservation comparable to "
        "professional translation as measured by automated metrics.",
    )

    add_heading(doc, "Study Design", level=2)
    add_paragraph(
        doc,
        "We conducted a cross-sectional evaluation of four frontier LLMs for medical document "
        "translation across eight languages, using automated metrics and validation against "
        "professional translation baselines.",
    )

    add_heading(doc, "Models and Languages", level=2)
    add_paragraph(
        doc,
        "The four models evaluated were GPT-5.1 (OpenAI, 2025), Claude Opus 4.5 (Anthropic, "
        "2025), Gemini 3 Pro (Google, 2025), and Kimi K2 Thinking (Moonshot AI, 2025). All "
        "models were accessed via their respective APIs using default parameters "
        "(temperature = 0.3). Models were evaluated on a single date using fixed API versions. "
        "We note that API-deployed models may be updated silently; exact version identifiers "
        "were recorded where accessible and are reported in Supplementary Table S0.",
    )
    add_paragraph(
        doc,
        "The eight target languages were Russian, Chinese (Simplified), Spanish, Vietnamese, "
        "Korean, Arabic, Tagalog, and Haitian Creole, representing diverse linguistic families, "
        "scripts, and resource levels. Arabic was evaluated using Modern Standard Arabic (MSA) "
        "prompting. Dialectal variation (Moroccan Darija, Levantine, Egyptian Arabic) was not "
        "assessed and may yield substantially different results; this represents a meaningful "
        "limitation for generalizability to Arabic-speaking patient populations in the "
        "United States.",
    )

    add_heading(doc, "Data Sample", level=2)
    add_paragraph(
        doc,
        "We assembled 22 medical education documents from two domains. Vaccine Information "
        "Statements (n = 11) were CDC documents distributed via Immunize.org covering "
        "Hepatitis B, HPV, Influenza, MMR, Meningococcal ACWY, Pneumococcal vaccines, Polio, "
        "Shingles, Tdap, and Varicella.(20) Cancer education materials (n = 11) were American "
        "Cancer Society documents covering post-diagnosis guidance (breast, cervical, colorectal, "
        "lung, prostate cancer), treatment side effects, and skin cancer resources.(21) Documents "
        "were selected based on the availability of professional translations in all eight target "
        "languages. We note that both CDC Vaccine Information Statements and ACS cancer education "
        "materials represent among the most carefully edited, ambiguity-minimized, and "
        "professionally reviewed documents in existence. Results may not generalize to clinical "
        "notes, consent forms, or physician-generated text, which contain substantially more "
        "ambiguity and domain-specific nuance.",
    )

    add_heading(doc, "Language Classification", level=2)
    add_paragraph(
        doc,
        "Languages were classified by CommonCrawl representation, a primary LLM training "
        "corpus:(22,23) high-resource (greater than 1%: Spanish, Chinese, Russian, Vietnamese), "
        "medium-resource (0.1-1%: Korean, Arabic), and low-resource (less than 0.1%: Tagalog, "
        "Haitian Creole). See Supplementary Table S1 for detailed classification criteria.",
    )

    add_heading(doc, "Translation Pipeline", level=2)
    add_paragraph(
        doc,
        "For each document-language-model combination, the pipeline proceeded as follows. First, "
        "forward translation converted the English source text to the target language using the "
        "LLM. Second, back-translation converted the target-language output back to English "
        "using the same LLM, with no access to the original English text. Third, professional "
        "back-translation converted the professionally-translated version of the same document "
        "back to English using the same LLM, serving as a methodological baseline. This yielded "
        "704 translation pairs (22 documents x 8 languages x 4 models). Two Kimi K2 translations "
        "were blocked by content filtering; the filtered documents involved vaccine-related "
        "content, and the specific filtering trigger was not disclosed by the API. These two "
        "pairs were excluded from analyses involving Kimi K2, leaving 702 evaluable pairs.",
    )

    add_heading(doc, "Cross-Model Sensitivity Analysis", level=2)
    add_paragraph(
        doc,
        "To address the concern that using the same model for forward and back-translation "
        "validates internal consistency rather than translation quality, we conducted a "
        "pre-specified cross-model back-translation sensitivity analysis. For each of the 702 "
        "successfully translated documents, we used each of the other three models (excluding "
        "Kimi K2 as back-translator due to its substantially higher latency of 60-90 minutes "
        "per document) to perform back-translation of the existing forward translations. Kimi K2 "
        "was excluded only as a back-translator; its forward translations were evaluated by all "
        "three other models. This yielded 1,576 additional cross-model back-translations across "
        "nine model pairings. We then computed LaBSE scores comparing the cross-model "
        "back-translations to the original English and compared these to the corresponding "
        "same-model LaBSE scores.",
    )

    add_heading(doc, "Evaluation Metrics", level=2)
    add_paragraph(
        doc,
        "Translation quality was assessed through two complementary lenses. Lexical fidelity "
        "measures whether the same words and phrases are used. Semantic fidelity measures "
        "whether meaning is preserved regardless of exact wording. High lexical fidelity implies "
        "high semantic fidelity, but the reverse is not true: a translation can preserve meaning "
        "while using different vocabulary. This distinction is important for cross-lingual "
        "evaluation, as languages with different scripts or morphological structures will "
        "inherently show lower lexical overlap even when meaning is fully preserved.",
    )
    add_paragraph(
        doc,
        "We used LaBSE (semantic similarity, 0-1 scale)(15) and BLEU (lexical overlap, "
        "0-100 scale)(24) to measure back-translation fidelity relative to the original English. "
        "For comparison to professional translations, we used BLEU, chrF,(27) BERTScore,(16) "
        "and COMET.(17) We note that COMET was developed and validated for neural machine "
        "translation systems and has not been formally validated for LLM-generated translations; "
        "BERTScore and LaBSE, as embedding-based metrics, do not carry this constraint. LaBSE "
        "scores greater than 0.90 indicate high semantic preservation. We acknowledge that the "
        "clinical meaningfulness of specific metric differences has not been established; "
        "anchoring to patient comprehension or clinical error rates would require human outcome "
        "studies beyond the scope of this evaluation.",
    )

    add_heading(doc, "Statistical Analysis", level=2)
    add_paragraph(
        doc,
        "We report means and standard deviations. Model comparisons used Kruskal-Wallis tests "
        "with Dunn's post-hoc correction.(28) Language comparisons were stratified by resource "
        "level. Significance was set at p less than 0.05. Analyses used Python 3.11 with "
        "SciPy (v1.13.1) and statsmodels (v0.14.1).",
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # RESULTS
    # -------------------------------------------------------------------------
    add_heading(doc, "Results", level=1)

    add_paragraph(
        doc,
        "We completed 702 of 704 translation pairs (99.7%). Two Kimi K2 translations failed "
        "due to content filtering on vaccine-related content.",
    )

    add_heading(doc, "Co-Primary Outcome 1: LLM Back-Translation Fidelity", level=2)
    add_paragraph(
        doc,
        "All models achieved high semantic preservation (Table 1). Kruskal-Wallis testing "
        "revealed significant differences between models for both LaBSE (H = 156.67, "
        "p less than 0.001) and BLEU (H = 170.69, p less than 0.001). Dunn's post-hoc tests "
        "with Bonferroni correction showed Claude Opus 4.5 significantly outperformed all other "
        "models on semantic preservation (all pairwise p less than 0.001). GPT-5.1 significantly "
        "outperformed Kimi K2 (p = 0.003) and Gemini 3 Pro (p less than 0.001), while Gemini 3 "
        "Pro and Kimi K2 showed no significant difference (p = 1.0).",
    )

    p = doc.add_paragraph()
    r = p.add_run("[Table 1 about here]")
    r.italic = True

    # Table 1
    add_heading(doc, "Table 1. Back-Translation Fidelity by Model", level=3)
    add_table(
        doc,
        headers=["Model", "n", "LaBSE (Mean +/- SD)", "BLEU (Mean +/- SD)"],
        rows=[
            ["Claude Opus 4.5", "176", "0.987 +/- 0.013", "68.7 +/- 8.3"],
            ["GPT-5.1", "176", "0.957 +/- 0.045", "64.3 +/- 7.4"],
            ["Kimi K2", "174", "0.940 +/- 0.053", "54.9 +/- 8.7"],
            ["Gemini 3 Pro", "176", "0.921 +/- 0.065", "61.5 +/- 9.3"],
        ],
        footer_note=(
            "Kruskal-Wallis: LaBSE H = 156.67, p < 0.001; BLEU H = 170.69, p < 0.001."
        ),
    )

    add_heading(doc, "Co-Primary Outcome 2: Comparison to Professional Translation", level=2)
    add_paragraph(
        doc,
        "LLM translations approached professional quality across all models (Table 2). "
        "Kruskal-Wallis testing showed no significant differences between models for "
        "COMET (H = 2.07, p = 0.56) or BLEU (H = 5.50, p = 0.14). BERTScore showed a "
        "significant effect (H = 13.58, p = 0.004), with post-hoc tests revealing Claude "
        "Opus 4.5 significantly outperformed GPT-5.1 (p = 0.04) and Kimi K2 (p = 0.004). "
        "COMET scores of 0.87-0.88 across models indicate performance in the range observed "
        "for professional translation systems; however, the clinical significance of differences "
        "within this range has not been established.",
    )

    p = doc.add_paragraph()
    r = p.add_run("[Table 2 about here]")
    r.italic = True

    # Table 2
    add_heading(doc, "Table 2. LLM vs Professional Translation Quality", level=3)
    add_table(
        doc,
        headers=["Model", "n", "BLEU", "BERTScore", "COMET"],
        rows=[
            ["Gemini 3 Pro", "176", "39.4 +/- 14.9", "0.845 +/- 0.060", "0.876 +/- 0.054"],
            ["Claude Opus 4.5", "176", "37.3 +/- 14.4", "0.859 +/- 0.043", "0.874 +/- 0.054"],
            ["GPT-5.1", "176", "36.0 +/- 13.8", "0.844 +/- 0.052", "0.871 +/- 0.052"],
            ["Kimi K2", "174", "36.0 +/- 14.3", "0.840 +/- 0.052", "0.872 +/- 0.052"],
        ],
        footer_note=(
            "Kruskal-Wallis: COMET H = 2.07, p = 0.56; BLEU H = 5.50, p = 0.14; "
            "BERTScore H = 13.58, p = 0.004."
        ),
    )

    add_heading(doc, "Cross-Model Sensitivity Analysis", level=2)
    add_paragraph(
        doc,
        "Across 1,550 matched pairs with both same-model and cross-model LaBSE scores, mean "
        "LaBSE was 0.9497 for same-model back-translation and 0.9488 for cross-model "
        "back-translation (mean delta: -0.0009). No model showed the substantial decline in "
        "fidelity expected if same-model circularity were driving results (Table 3). Notably, "
        "Kimi K2 and Gemini 3 Pro both showed marginally higher scores in cross-model than "
        "same-model conditions, indicating that their primary study scores were not inflated "
        "by circularity. Low-resource languages scored 0.9482 in cross-model conditions compared "
        "to 0.9454 for high-resource languages, with the direction of any difference favoring "
        "low-resource languages (p = 0.025 by Mann-Whitney U, favoring low-resource).",
    )

    p = doc.add_paragraph()
    r = p.add_run("[Table 3 about here]")
    r.italic = True

    # Table 3
    add_heading(
        doc,
        "Table 3. Cross-Model Sensitivity Analysis: LaBSE by Model Pairing",
        level=3,
    )
    add_table(
        doc,
        headers=[
            "Forward Model",
            "Back-Translation Model",
            "LaBSE (Cross-Model)",
            "Same-Model LaBSE",
            "Delta",
        ],
        rows=[
            ["Claude Opus 4.5", "GPT-5.1", "0.9737", "0.9869", "-0.0132"],
            ["Claude Opus 4.5", "Gemini 3 Pro", "0.9784", "0.9869", "-0.0085"],
            ["GPT-5.1", "Claude Opus 4.5", "0.9596", "0.9564", "+0.0032"],
            ["GPT-5.1", "Gemini 3 Pro", "0.9498", "0.9564", "-0.0066"],
            ["Kimi K2", "GPT-5.1", "0.9452", "0.9398", "+0.0054"],
            ["Kimi K2", "Claude Opus 4.5", "0.9434", "0.9398", "+0.0036"],
            ["Kimi K2", "Gemini 3 Pro", "0.9414", "0.9398", "+0.0016"],
            ["Gemini 3 Pro", "GPT-5.1", "0.9245", "0.9214", "+0.0031"],
            ["Gemini 3 Pro", "Claude Opus 4.5", "0.9242", "0.9214", "+0.0028"],
        ],
        footer_note=(
            "Overall: same-model LaBSE mean 0.9497; cross-model LaBSE mean 0.9488; "
            "mean delta -0.0009 (n = 1,550 matched pairs). Kimi K2 was excluded as a "
            "back-translation model due to high latency (60-90 minutes per document); "
            "its forward translations were evaluated by all three other models."
        ),
    )

    add_heading(doc, "Language Performance by Resource Level", level=2)
    add_paragraph(
        doc,
        "Low-resource languages achieved semantic preservation that did not differ significantly "
        "from high-resource languages (Figure 2, Supplementary Table S3). Tagalog (0.950) and "
        "Haitian Creole (0.955) scored comparably to Spanish (0.954) and Vietnamese (0.953). "
        "Mann-Whitney U testing found no significant difference between resource groups for "
        "semantic preservation (LaBSE: 0.948 vs 0.952, p = 0.066). This result represents the "
        "absence of a detected difference, not demonstrated equivalence; the study was not "
        "powered to exclude small but potentially meaningful differences.",
    )
    add_paragraph(
        doc,
        "BLEU scores varied substantially by language (range: 15.5-54.3), while semantic "
        "metrics remained consistently high across all languages (LaBSE range: 0.937-0.976). "
        "High-resource languages showed better alignment with professional translations "
        "(COMET: 0.879 vs 0.837, p less than 0.001), indicating differences in phrasing "
        "convention even where semantic meaning was preserved.",
    )

    p = doc.add_paragraph()
    r = p.add_run("[Figure 2 about here]")
    r.italic = True

    add_heading(doc, "Document Category", level=2)
    add_paragraph(
        doc,
        "Cancer education materials achieved substantially higher back-translation fidelity "
        "(mean LaBSE: 0.984) than Vaccine Information Statements (mean LaBSE: 0.919; "
        "Mann-Whitney U, p less than 0.001). This gap likely reflects differences in document "
        "structure and linguistic complexity: cancer materials tend to use more descriptive prose "
        "with concrete medical facts, while vaccine information statements combine descriptive "
        "content with procedural safety language, conditional statements, and tabular "
        "side-effect listings. The procedural and conditional language in vaccine documents may "
        "be more susceptible to meaning drift in translation. This difference is worth attending "
        "to in clinical deployment decisions about which document types are most suitable for "
        "LLM translation without human review.",
    )

    add_heading(doc, "Methodological Validation", level=2)
    add_paragraph(
        doc,
        "Back-translation of professional translations yielded high fidelity with original "
        "English (LaBSE: 0.92-0.94; BLEU: 43.9-55.8; Supplementary Table S2), establishing "
        "a benchmark context for LLM performance.",
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # DISCUSSION
    # -------------------------------------------------------------------------
    add_heading(doc, "Discussion", level=1)

    add_paragraph(
        doc,
        "This study evaluated LLM translation fidelity for medical patient education materials "
        "across eight languages using two co-primary outcomes: back-translation consistency and "
        "comparison to professional translation. Across 702 translation pairs, all frontier "
        "models achieved high semantic preservation (LaBSE greater than 0.92) and approached "
        "professional translation quality on COMET (0.87-0.88). A cross-model sensitivity "
        "analysis involving 1,576 additional back-translations across nine model pairings found "
        "a mean LaBSE difference of -0.0009 compared to same-model back-translation, "
        "substantially mitigating concerns about circularity in the primary outcome. Low-resource "
        "languages (Tagalog, Haitian Creole) did not perform significantly worse than "
        "high-resource languages, though this finding requires careful interpretation given "
        "methodological constraints discussed below.",
    )

    add_heading(doc, "Addressing the Circularity Concern", level=2)
    add_paragraph(
        doc,
        "A fundamental methodological question in back-translation studies is whether using the "
        "same model for forward and back-translation tests translation quality or merely internal "
        "consistency. A model with systematic, semantically coherent errors in both directions "
        "could score near-perfectly while producing reliably wrong translations. Our cross-model "
        "sensitivity analysis directly tests this concern. Using each model's existing forward "
        "translations and back-translating with three independently-developed models (GPT-5.1, "
        "Claude Opus 4.5, Gemini 3 Pro, representing three distinct organizational and "
        "architectural lineages), we found no evidence of score inflation attributable to "
        "same-model consistency. The mean LaBSE delta was -0.0009 across 1,550 matched pairs. "
        "Kimi K2 and Gemini 3 Pro showed marginally higher cross-model than same-model scores, "
        "directly contrary to what circularity would predict. The professional translation "
        "comparison (co-primary outcome 2) provides a fully independent validation pathway with "
        "no circularity: it compares LLM translations to human-generated reference translations "
        "without any back-translation step.",
    )

    add_heading(doc, "Interpreting the Low-Resource Language Finding", level=2)
    add_paragraph(
        doc,
        "The finding that Tagalog and Haitian Creole achieved LaBSE scores statistically "
        "indistinguishable from high-resource languages (p = 0.066) warrants cautious "
        "interpretation on two grounds.",
    )
    add_paragraph(
        doc,
        "First, p = 0.066 represents absence of detected difference, not demonstrated "
        "equivalence. The study was not powered to exclude clinically meaningful differences; "
        "the observed difference (LaBSE: 0.948 vs 0.952) might be practically irrelevant or "
        "might matter depending on thresholds that patient comprehension studies have yet to "
        "establish.",
    )
    add_paragraph(
        doc,
        "Second, both Tagalog and Haitian Creole exhibit substantial lexical borrowing from "
        "English and French, respectively, in medical domains. Medical Tagalog frequently "
        "retains English terminology (e.g., \"chemotherapy,\" \"diabetes,\" \"vaccine\"), as do "
        "many post-colonial languages in technical contexts. Similarly, Haitian Creole "
        "incorporates substantial French and English vocabulary in medical settings.(30) This "
        "lexical overlap could inflate back-translation fidelity scores relative to languages "
        "such as Chinese or Russian, which use native medical terminology that requires genuine "
        "translation. In effect, the model may be routing English medical terms through a "
        "foreign-language shell and returning them largely intact, producing high LaBSE scores "
        "without having translated the medically critical content. We regard this as a genuine "
        "threat to the primary interpretation rather than a peripheral caveat. Future work "
        "should specifically examine whether performance differs systematically between borrowed "
        "and native medical terminology within Tagalog and Haitian Creole.",
    )
    add_paragraph(
        doc,
        "Importantly, however, this concern is not resolved by the cross-model analysis, which "
        "would be equally susceptible to the same transliteration pattern. The professional "
        "translation comparison remains the most independent check: it examines whether "
        "LLM-generated translations resemble what human professional translators produce, "
        "regardless of how back-translation scores are generated.",
    )

    add_heading(doc, "Professional Translation Comparison as Co-Primary Outcome", level=2)
    add_paragraph(
        doc,
        "COMET scores of 0.87-0.88 across models indicate performance in the range of "
        "professional-quality machine translation systems on this metric. This result does not "
        "depend on back-translation and thus avoids the circularity concern entirely. The absence "
        "of significant between-model differences on COMET (p = 0.56) suggests that at the "
        "frontier level, model selection matters less than language pair and document type for "
        "this class of materials. The clinical significance of differences within this COMET "
        "range has not been established. We report this finding descriptively and encourage "
        "future work to anchor COMET differences to patient comprehension outcomes and clinical "
        "error rates.",
    )

    add_heading(doc, "Document Type and Generalizability", level=2)
    add_paragraph(
        doc,
        "The 65-point gap in back-translation LaBSE between cancer materials (0.984) and vaccine "
        "documents (0.919) warrants attention. CDC Vaccine Information Statements and ACS cancer "
        "education materials are among the most carefully edited, ambiguity-minimized documents "
        "produced in healthcare. They represent an upper bound on what LLM translation might "
        "reliably achieve. Generalizability to discharge instructions, consent forms, clinical "
        "notes, or other documents with greater ambiguity, conditional language, or "
        "physician-generated content cannot be assumed from these results. The vaccine-specific "
        "gap suggests that documents with denser procedural and conditional content are more "
        "susceptible to meaning drift, which should inform decisions about which document types "
        "are prioritized for human review.",
    )

    add_heading(doc, "Arabic Dialect Limitation", level=2)
    add_paragraph(
        doc,
        "Arabic was evaluated using Modern Standard Arabic prompting. In the United States, "
        "Arabic-speaking communities include speakers of Moroccan Darija, Levantine, Egyptian, "
        "and Gulf dialects, which differ substantially from MSA in vocabulary, syntax, and "
        "comprehensibility for typical patients. The high mean LaBSE for Arabic (0.976 in the "
        "primary analysis; 0.979 in the cross-model analysis) likely reflects MSA performance "
        "specifically and should not be generalized to dialectal Arabic speakers, who represent "
        "the majority of Arabic-speaking patients in U.S. healthcare settings.",
    )

    add_heading(doc, "Limitations", level=2)
    add_paragraph(
        doc,
        "In addition to the methodological considerations discussed above, the following "
        "limitations apply. First, all evaluation was conducted using automated metrics. While "
        "multiple complementary metrics were employed, they cannot fully capture clinical "
        "accuracy, cultural appropriateness, or patient comprehension. Human evaluation with "
        "bilingual clinicians or patients remains the standard for definitively assessing "
        "translation quality. Second, this evaluation represents a snapshot of model capabilities "
        "at a specific point in late 2025. API-deployed models may be updated silently, limiting "
        "reproducibility. Third, only English-to-target translation was evaluated; "
        "target-to-English translation may show different patterns. Fourth, the document corpus "
        "was limited to two types of standardized patient education materials, limiting "
        "generalizability. Fifth, no patient comprehension or clinical error rate data were "
        "collected; the study cannot support conclusions about clinical safety or patient "
        "outcomes.",
    )

    add_heading(doc, "Clinical and Policy Implications", level=2)
    add_paragraph(
        doc,
        "These findings suggest that frontier LLMs show promise for extending medical "
        "translation access to speakers of historically underserved languages, particularly for "
        "standardized patient education materials. The convergent evidence from two co-primary "
        "outcomes and an independent cross-model sensitivity analysis strengthens confidence in "
        "the primary findings compared to single-metric evaluations.",
    )
    add_paragraph(
        doc,
        "We emphasize that these results do not, on their own, support relaxing existing human "
        "review requirements for machine translation in clinical settings. Section 1557 of the "
        "Affordable Care Act and its implementing regulations require qualified human translator "
        "review of critical translated materials, and this study did not assess clinical error "
        "rates, patient comprehension, or harm outcomes. The appropriate role for evidence of "
        "this kind is to inform the design of future studies with human outcome measures, not to "
        "directly guide regulatory decisions. We encourage prospective work that links automated "
        "metric performance to patient-facing outcomes to establish what metric thresholds are "
        "clinically meaningful.",
    )

    add_heading(doc, "Future Directions", level=2)
    add_paragraph(
        doc,
        "Several extensions of this work are warranted. First, evaluation should expand to "
        "additional clinical document types, including clinician notes, medication instructions, "
        "and informed consent forms, to better characterize the boundary of reliable LLM "
        "translation. Second, given the rapid evolution of LLMs, longitudinal assessment is "
        "needed to track performance as models are updated. Third, human comprehension studies "
        "should be conducted in partnership with patients across target language communities, "
        "including dialect speakers within Arabic, Spanish, and Chinese communities. Fourth, the "
        "specific contribution of lexical borrowing to Tagalog and Haitian Creole scores should "
        "be quantified by comparing performance on borrowed-term-heavy versus "
        "native-terminology-heavy passages. Fifth, future regulatory frameworks should consider "
        "whether performance-based criteria, anchored to human comprehension outcomes, might "
        "complement or refine existing categorical requirements in a manner that preserves "
        "patient safety while recognizing demonstrated technological capabilities.",
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # CONCLUSION
    # -------------------------------------------------------------------------
    add_heading(doc, "Conclusion", level=1)
    add_paragraph(
        doc,
        "Frontier LLMs preserve medical meaning in standardized patient education materials "
        "across both high-resource and low-resource languages, with performance approaching "
        "professional translation on automated metrics. A cross-model sensitivity analysis "
        "involving 1,576 back-translations across nine model pairings found no evidence of score "
        "inflation attributable to same-model circularity. Low-resource languages did not "
        "perform significantly worse than high-resource languages, though this finding requires "
        "further study with patient comprehension outcomes and human evaluation before informing "
        "clinical or regulatory decisions.",
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # DATA SHARING STATEMENT
    # -------------------------------------------------------------------------
    add_heading(doc, "Data Sharing Statement", level=1)
    add_paragraph(
        doc,
        "All source documents, translation outputs, and evaluation metrics are available at: "
        "https://github.com/hallixrap/llm-translation-study",
    )

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # REFERENCES
    # -------------------------------------------------------------------------
    add_heading(doc, "References", level=1)

    references = [
        "1. Simonsmeier BA, Flaig M, Simacek T, Schneider M. What sixty years of research says "
        "about the effectiveness of patient education on health: a second order meta-analysis. "
        "Health Psychol Rev. 2022;16(3):450-474.",

        "2. U.S. Census Bureau. American Community Survey 2023 1-Year Estimates, Table B16001: "
        "Language Spoken at Home by Ability to Speak English for the Population 5 Years and "
        "Over. U.S. Census Bureau; 2024.",

        "3. Twersky SE, Jefferson R, Garcia-Ortiz L. The impact of limited English proficiency "
        "on healthcare access and outcomes in the U.S.: a scoping review. Healthcare. "
        "2024;12(3):364.",

        "4. Divi C, Koss RG, Schmaltz SP, Loeb JM. Language proficiency and adverse events in "
        "US hospitals: a pilot study. Int J Qual Health Care. 2007;19(2):60-67.",

        "5. The Joint Commission. Advancing Effective Communication, Cultural Competence, and "
        "Patient- and Family-Centered Care: A Roadmap for Hospitals. 2010.",

        "6. Department of Health and Human Services. Nondiscrimination in Health Programs and "
        "Activities. Fed Regist. 2024;89(88):37522-37703.",

        "7. Office of Minority Health. National Standards for Culturally and Linguistically "
        "Appropriate Services (CLAS) in Health and Health Care. U.S. Department of Health and "
        "Human Services; 2013.",

        "8. Davis SH, Rosenberg J, Nguyen J, et al. Translating discharge instructions for "
        "limited English-proficient families: strategies and barriers. Hosp Pediatr. "
        "2019;9(10):779-787.",

        "9. Lopez I, Velasquez DE, Chen JH, Rodriguez JA. Operationalizing machine-assisted "
        "translation in healthcare. npj Digit Med. 2025;8:584.",

        "10. Choe AY, Schondelmeyer AC, Thomson J, et al. Improving discharge instructions for "
        "hospitalized children with limited English proficiency. Hosp Pediatr. "
        "2021;11(11):1213-1222.",

        "11. Brown T, Mann B, Ryder N, et al. Language models are few-shot learners. Advances "
        "in Neural Information Processing Systems. 2020;33:1877-1901.",

        "12. Quan K, Lynch J. The High Costs of Language Barriers in Medical Malpractice. "
        "National Health Law Program; 2010.",

        "13. Brewster RC, Gonzalez P, Khazanchi R, et al. Performance of ChatGPT and Google "
        "Translate for pediatric discharge instruction translation. Pediatrics. "
        "2024;154(1):e2023065573.",

        "14. American Medical Association. Physician sentiments around the use of AI in health "
        "care: motivations, opportunities, risks, and use cases. Published February 2025.",

        "15. Feng F, Yang Y, Cer D, et al. Language-agnostic BERT sentence embedding. "
        "Proceedings of ACL. 2022:878-891.",

        "16. Zhang T, Kishore V, Wu F, Weinberger KQ, Artzi Y. BERTScore: Evaluating text "
        "generation with BERT. Proceedings of ICLR. 2020.",

        "17. Rei R, Stewart C, Farinha AC, Lavie A. COMET: A neural framework for MT "
        "evaluation. Proceedings of EMNLP. 2020:2685-2702.",

        "18. Brislin RW. Back-translation for cross-cultural research. Journal of "
        "Cross-Cultural Psychology. 1970;1(3):185-216.",

        "19. Tyupa S. A theoretical framework for back-translation as a quality assessment tool. "
        "New Voices in Translation Studies. 2011;7(1):35-46.",

        "20. Immunize.org. Vaccine Information Statements (VIS). "
        "https://www.immunize.org/vis/. Accessed December 20, 2025.",

        "21. American Cancer Society. Cancer Information. "
        "https://www.cancer.org/cancer.html. Accessed December 20, 2025.",

        "22. Lai VD, Ngo NT, Pouran Ben Veyseh A, et al. ChatGPT beyond English: towards a "
        "comprehensive evaluation of large language models in multilingual learning. "
        "Findings of EMNLP 2023. 2023:13171-99.",

        "23. Common Crawl. Statistics of Common Crawl Monthly Archives: Languages. "
        "https://commoncrawl.github.io/cc-crawl-statistics/plots/languages.html. "
        "Accessed December 20, 2025.",

        "24. Papineni K, Roukos S, Ward T, Zhu WJ. BLEU: a method for automatic evaluation "
        "of machine translation. Proceedings of ACL. 2002:311-318.",

        "25. Conneau A, Khandelwal K, Goyal N, et al. Unsupervised cross-lingual representation "
        "learning at scale. Proceedings of ACL. 2020:8440-8451.",

        "26. Devlin J, Chang MW, Lee K, Toutanova K. BERT: Pre-training of deep bidirectional "
        "transformers for language understanding. Proceedings of NAACL. 2019:4171-4186.",

        "27. Popovic M. chrF: character n-gram F-score for automatic MT evaluation. "
        "Proceedings of WMT. 2015:392-395.",

        "28. Dunn OJ. Multiple comparisons using rank sums. Technometrics. 1964;6(3):241-252.",

        "29. Office for Civil Rights, U.S. Department of Health and Human Services. Dear "
        "Colleague Letter: Language Access and Section 1557 of the Affordable Care Act. "
        "October 2024.",

        "30. Baklanova E. Types of borrowings in Tagalog/Filipino. Kritika Kultura. "
        "2017;28:35-54.",

        "31. Silva MS, Ahmadi S. Language models are borrowing-blind: a multilingual evaluation "
        "of loanword identification across 10 languages. arXiv preprint arXiv:2510.26254. 2025.",
    ]

    for ref in references:
        add_paragraph(doc, ref)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # ACKNOWLEDGMENTS
    # -------------------------------------------------------------------------
    add_heading(doc, "Acknowledgments", level=1)
    add_paragraph(
        doc,
        "We thank the CDC, Immunize.org, and the American Cancer Society for making "
        "professionally-translated health education materials publicly available.",
    )

    # -------------------------------------------------------------------------
    # AUTHOR CONTRIBUTIONS
    # -------------------------------------------------------------------------
    add_heading(doc, "Author Contributions", level=1)
    contributions = [
        ("Concept and design: ", "Anyaegbuna, Ma, Steele, Liang, Chen, Schulman"),
        (
            "Acquisition, analysis, or interpretation of data: ",
            "Anyaegbuna, Ma, Liang, Steele, Perez Guerrero, Liu",
        ),
        ("Drafting of the manuscript: ", "Anyaegbuna"),
        (
            "Critical revision of the manuscript for important intellectual content: ",
            "All authors",
        ),
        ("Statistical analysis: ", "Anyaegbuna, Keyes"),
        ("Administrative, technical, or material support: ", "All authors"),
        ("Supervision: ", "Chen, Schulman"),
    ]
    for label, text in contributions:
        add_bold_label_paragraph(doc, label, text)

    # -------------------------------------------------------------------------
    # FUNDING
    # -------------------------------------------------------------------------
    add_heading(doc, "Funding", level=1)
    add_paragraph(
        doc,
        "Support for this research was provided by the Commonwealth Fund. The views presented "
        "here are those of the authors and should not be attributed to the Commonwealth Fund "
        "or its directors, officers, or staff.",
    )

    # -------------------------------------------------------------------------
    # CONFLICTS OF INTEREST
    # -------------------------------------------------------------------------
    add_heading(doc, "Conflicts of Interest", level=1)
    add_paragraph(doc, "The authors declare no conflicts of interest.")

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # SUPPLEMENTARY TABLES
    # -------------------------------------------------------------------------
    add_heading(doc, "Supplementary Tables", level=1)

    add_heading(doc, "Supplementary Table S1: Language Classification Details", level=2)
    add_table(
        doc,
        headers=["Language", "Resource Level", "CommonCrawl %", "U.S. Speakers (millions)"],
        rows=[
            ["Russian", "High", "6.48", "0.9"],
            ["Chinese (Simplified)", "High", "6.18", "3.5"],
            ["Spanish", "High", "4.41", "41.8"],
            ["Vietnamese", "High", "1.08", "1.6"],
            ["Korean", "Medium", "0.80", "1.1"],
            ["Arabic", "Medium", "0.67", "1.3"],
            ["Tagalog", "Low", "0.008", "1.8"],
            ["Haitian Creole", "Low", "0.003", "0.9"],
        ],
        footer_note=(
            "High-resource: >1%; medium-resource: 0.1-1%; low-resource: <0.1%. "
            "U.S. speaker populations from American Community Survey.(2)"
        ),
    )

    add_heading(doc, "Supplementary Table S2: Professional Translation Back-Translation Fidelity", level=2)
    add_table(
        doc,
        headers=["Model", "BLEU", "BERTScore", "LaBSE"],
        rows=[
            ["GPT-5.1", "55.8", "0.919", "0.940"],
            ["Claude Opus 4.5", "54.7", "0.924", "0.942"],
            ["Gemini 3 Pro", "48.3", "0.913", "0.928"],
            ["Kimi K2", "43.9", "0.912", "0.934"],
        ],
    )

    add_heading(doc, "Supplementary Table S3: Translation Fidelity by Language", level=2)
    add_table(
        doc,
        headers=["Language", "Resource Level", "Back-Translation LaBSE", "vs Professional BLEU"],
        rows=[
            ["Arabic", "Medium", "0.976", "41.6"],
            ["Haitian Creole", "Low", "0.955", "37.2"],
            ["Spanish", "High", "0.954", "54.3"],
            ["Vietnamese", "High", "0.953", "50.1"],
            ["Tagalog", "Low", "0.950", "43.8"],
            ["Russian", "High", "0.945", "33.4"],
            ["Chinese (Simplified)", "High", "0.942", "15.5"],
            ["Korean", "Medium", "0.937", "21.7"],
        ],
    )

    # -------------------------------------------------------------------------
    # SAVE
    # -------------------------------------------------------------------------
    doc.save(OUTPUT_PATH)
    print(f"Manuscript saved to:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build_manuscript()
